"""
Registration Manager Tools - CSV/Excel Processing, Sorting, Filtering, and Partitioning
"""

import csv
import datetime
import os
import re
from typing import Any, List, Tuple

try:
    from docx import Document
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    Document = None
    OxmlElement = parse_xml = nsdecls = qn = Inches = Pt = RGBColor = None

# BASE_DIR represents the absolute path of this agent's folder, ensuring self-contained integrations
BASE_DIR = os.path.dirname(os.path.abspath(__file__))


def resolve_path(rel_path: str) -> str:
    """Resolves a path relative to the agent's folder, falling back to CWD and workspace root."""
    if not rel_path:
        return ""
    if os.path.isabs(rel_path):
        return rel_path

    # 1. Direct relative to BASE_DIR
    local_path = os.path.abspath(os.path.join(BASE_DIR, rel_path))
    if os.path.exists(local_path):
        return local_path

    # 2. Strip leading 'agents/registration_manager/' or 'registration_manager/'
    stripped = re.sub(r"^(agents/)?registration_manager/", "", rel_path)
    stripped_path = os.path.abspath(os.path.join(BASE_DIR, stripped))
    if os.path.exists(stripped_path):
        return stripped_path

    # 3. Check relative to workspace root (2 levels up)
    workspace_root = os.path.abspath(os.path.join(BASE_DIR, "..", ".."))
    workspace_path = os.path.abspath(os.path.join(workspace_root, rel_path))
    if os.path.exists(workspace_path):
        return workspace_path

    workspace_agents_path = os.path.abspath(os.path.join(workspace_root, "agents", rel_path))
    if os.path.exists(workspace_agents_path):
        return workspace_agents_path

    return stripped_path


async def stage_uploaded_registration(file_path: str, tool_context: Any) -> str:
    """Finds and stages the active user-uploaded CSV or Excel registration file from the chat session events.
    If no chat attachment is found, falls back to the provided local file_path.

    Args:
        file_path: Fallback local file path (if no chat attachment is uploaded).
    """
    print("\n📥 [Tool: stage_uploaded_registration] Resolving user uploaded registration file from session...")

    # 1. ALWAYS scan session events first to check if the user uploaded a document in the active chat.
    try:
        session = tool_context.session
        events = session.events or []
        print(f"DEBUG: Staging scanning total events: {len(events)}")

        # Iterate events in REVERSE to find the most recent user-uploaded file
        for event in reversed(events):
            if event.author != "user":
                continue
            if not event.content or not event.content.parts:
                continue
            for part in event.content.parts:
                # Check for inline_data
                if part.inline_data and part.inline_data.data:
                    mime = part.inline_data.mime_type or ""
                    data = part.inline_data.data
                    print(f"DEBUG: Found user attachment with mime: {mime} ({len(data)} bytes)")

                    # Deduce extension based on MIME
                    ext = ".csv"
                    if "excel" in mime or "spreadsheet" in mime or "officedocument" in mime:
                        ext = ".xlsx"
                    elif "csv" in mime or "comma" in mime:
                        ext = ".csv"

                    save_path = os.path.join(BASE_DIR, "results", f"staged_registrations{ext}")
                    os.makedirs(os.path.dirname(save_path), exist_ok=True)

                    import base64

                    file_bytes = base64.b64decode(data)
                    with open(save_path, "wb") as f:
                        f.write(file_bytes)

                    print(f"✅ Successfully staged user uploaded registration file to: {save_path}")
                    return f"registration_manager/results/staged_registrations{ext}"
    except Exception as e:
        print(f"⚠️ Could not extract attachment from session events: {e}")

    # 2. Fallback to local file path
    print("⚠️ No chat attachment found. Checking local file path fallback...")
    resolved = resolve_path(file_path)
    if os.path.exists(resolved) and os.path.isfile(resolved):
        ext = os.path.splitext(resolved.lower())[1]
        save_path = os.path.join(BASE_DIR, "results", f"staged_registrations{ext}")
        os.makedirs(os.path.dirname(save_path), exist_ok=True)
        import shutil

        shutil.copy2(resolved, save_path)
        print(f"✅ Staged local registration file to: {save_path}")
        return f"registration_manager/results/staged_registrations{ext}"

    return (
        "Error: No registration file was found in session events or local path. "
        "Please upload a CSV or Excel file or paste the participant names directly in the chat."
    )


def stage_manual_text_registrations(text_content: str) -> str:
    """Takes a raw text block of names/registrants passed manually in the chat,
    parses it, writes it to a temporary CSV file, and returns the path to it.

    Args:
        text_content: Raw text content containing the list of registrants.
    """
    print("\n📝 [Tool: stage_manual_text_registrations] Staging manual text registration list...")
    lines = [line.strip() for line in text_content.strip().split("\n") if line.strip()]

    # Determine if it's already comma/semicolon/tab separated CSV-like text
    # Look at the first line
    has_delimiters = False
    delimiter = ","
    if lines:
        first_line = lines[0]
        for d in [",", ";", "\t"]:
            if len(first_line.split(d)) >= 2:
                has_delimiters = True
                delimiter = d
                break

    save_path = os.path.join(BASE_DIR, "results", "staged_manual_registrations.csv")
    os.makedirs(os.path.dirname(save_path), exist_ok=True)

    with open(save_path, "w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f)
        if has_delimiters:
            # It already looks like CSV/TSV, parse and write
            for line in lines:
                row = [col.strip() for col in line.split(delimiter)]
                writer.writerow(row)
        else:
            # It's a simple list of names (one per line)
            # We write a header "Full Name" and then each name
            writer.writerow(["Full Name"])
            for line in lines:
                writer.writerow([line])

    print(f"✅ Staged manual text registrations to: {save_path}")
    return "registration_manager/results/staged_manual_registrations.csv"


def parse_date(date_str: str) -> datetime.datetime:
    """Intelligently parses common date and timestamp formats into a datetime object for chronological sorting."""
    if not date_str:
        return datetime.datetime.max

    date_str = date_str.strip()

    # Try common formats
    formats = [
        "%Y-%m-%d %H:%M:%S",
        "%Y-%m-%dT%H:%M:%S",
        "%d.%m.%Y %H:%M:%S",
        "%d/%m/%Y %H:%M:%S",
        "%Y-%m-%d",
        "%d.%m.%Y",
        "%d/%m/%Y",
        "%I:%M:%S %p",
    ]

    for fmt in formats:
        try:
            dt = datetime.datetime.strptime(date_str, fmt)
            return dt.replace(tzinfo=None)
        except ValueError:
            continue

    # Try ISO date parsing or timestamp fallback
    try:
        dt = datetime.datetime.fromisoformat(date_str.replace("Z", "+00:00"))
        return dt.replace(tzinfo=None)
    except ValueError:
        pass

    # Standard fallback
    return datetime.datetime.max


def clean_and_validate_name(first_name: str, last_name: str) -> Tuple[str, str, str]:
    """Cleans names and validates if they are complete and clear.
    Returns: (cleaned_first_name, cleaned_last_name, reason_if_invalid)
    """
    fn = first_name.strip()
    ln = last_name.strip()

    if not fn and not ln:
        return "", "", "Empty name and surname"
    if not fn:
        return "", ln.capitalize(), "Missing first name"
    if not ln:
        return fn.capitalize(), "", "Missing last name"

    # Basic formatting
    fn_clean = " ".join(part.capitalize() for part in fn.split())
    ln_clean = " ".join(part.capitalize() for part in ln.split())

    # Check for invalid characters (symbols or digits)
    # We allow any alphabetic/letter characters (Unicode friendly) plus spaces, hyphens, dots, and apostrophes
    def is_valid_name_str(s: str) -> bool:
        return all(char.isalpha() or char.isspace() or char in "-.'" for char in s)

    if not is_valid_name_str(fn_clean) or not is_valid_name_str(ln_clean):
        return fn_clean, ln_clean, "Contains invalid characters, numbers, or symbols"

    # Check for short or incomplete names (e.g. single letters or abbreviations)
    # Allows valid short names like "Li" or "An" but flags single characters without dots like "J"
    if (len(fn_clean.replace(".", "")) <= 1 and not fn_clean.endswith(".")) or (
        len(ln_clean.replace(".", "")) <= 1 and not ln_clean.endswith(".")
    ):
        return fn_clean, ln_clean, "Incomplete name or initials without proper format"

    # Check for test keywords
    test_keywords = {"test", "demo", "asd", "qwerty", "guest", "anonymous"}
    fn_lower = fn_clean.lower()
    ln_lower = ln_clean.lower()
    if any(kw in fn_lower or kw in ln_lower for kw in test_keywords):
        return fn_clean, ln_clean, "Test or anonymous registration name"

    return fn_clean, ln_clean, ""


def get_script_priority(first_name: str, last_name: str) -> Tuple[int, str]:
    """Returns a script priority and group name:
    (1, "Latin") -> English, Polish, French, German, etc.
    (2, "Cyrillic") -> Ukrainian, Belarusian, Russian, Bulgarian, etc.
    (3, "Other") -> Other languages (Greek, Hebrew, Arabic, Asian scripts, etc.)
    """
    text = (first_name + last_name).replace(" ", "")
    if not text:
        return 3, "Other"

    # Cyrillic check: Unicode block U+0400 to U+04FF
    if any("\u0400" <= char <= "\u04ff" for char in text):
        return 2, "Cyrillic"

    # Check non-Latin unicode ranges
    # Greek: U+0370 to U+03FF
    # Hebrew: U+0590 to U+05FF
    # Arabic: U+0600 to U+06FF
    # CJK (Asian): U+4E00 to U+9FFF
    non_latin = False
    for char in text:
        val = ord(char)
        if (
            (0x0370 <= val <= 0x03FF)
            or (0x0590 <= val <= 0x05FF)
            or (0x0600 <= val <= 0x06FF)
            or (0x4E00 <= val <= 0x9FFF)
        ):
            non_latin = True
            break

    if non_latin:
        return 3, "Other"

    return 1, "Latin"


def load_organisers() -> List[str]:
    """Loads GDG organisers list from organisers.txt file or returns default list if not found."""
    # Try the root configs/ folder first, then the agent's folder, then fall back to defaults
    parent_dir = os.path.dirname(BASE_DIR)
    project_root = os.path.dirname(parent_dir)
    org_file = None
    for path in [
        os.path.join(project_root, "configs", "organisers.txt"),
        os.path.join(parent_dir, "configs", "organisers.txt"),
        os.path.join(BASE_DIR, "organisers.txt"),
    ]:
        if os.path.exists(path):
            org_file = path
            break

    if not org_file:
        return [
            "Aliaksandr Aplachykau",
            "Yuliya Aplachykava",
            "Aryna Stsiapanava",
            "Roman Hordiichuk",
            "Paul Kastel",
            "Rafał Płonka",
            "Yusuf Gültaç",
        ]
    try:
        with open(org_file, "r", encoding="utf-8") as f:
            return [line.strip() for line in f if line.strip()]
    except Exception as e:
        print(f"⚠️ Error loading organisers file: {e}")
        return [
            "Aliaksandr Aplachykau",
            "Yuliya Aplachykava",
            "Aryna Stsiapanava",
            "Roman Hordiichuk",
            "Paul Kastel",
            "Rafał Płonka",
            "Yusuf Gültaç",
        ]


def parse_manual_name_line(line: str) -> Tuple[str, str, str]:
    """Parses a manual name entry like 'John Doe (speaker)' or 'Jane Smith (organiser)' or 'Aliaksandr Aplachykau'.
    Returns: (first_name_clean, last_name_clean, status_override)
    """
    line = line.strip()
    status_override = ""

    # Try to extract status in parentheses
    match = re.search(r"\((speaker|organiser|organizer|registered|waitlist)\)", line, re.IGNORECASE)
    if match:
        status_override = match.group(1).lower()
        if status_override == "organizer":
            status_override = "organiser"
        line = re.sub(r"\(.*?\)", "", line).strip()

    # Split into first and last name
    parts = line.split(None, 1)
    fn = parts[0] if parts else ""
    ln = parts[1] if len(parts) > 1 else ""

    fn_clean, ln_clean, _ = clean_and_validate_name(fn, ln)
    return fn_clean, ln_clean, status_override


def process_registrations(file_path: str, capacity: int = 0, manual_confirmed: str = "") -> str:
    """Loads a CSV or Excel registration file, cleans, sorts (Latin vs Cyrillic vs Other scripts),
    handles capacity limits, manual overrides, and pre-existing waitlist marks.

    Args:
        file_path: Relative or absolute local file path of the CSV/XLSX file.
        capacity: Optional maximum attendee capacity. If 0 or None, defaults to pre-existing waitlist marks.
        manual_confirmed: Optional comma-separated or newline-separated list of names to force-add to Confirmed List.
    """
    print("\n📂 [Tool: process_registrations] Starting processing of registration list...")
    print(f"   ├─ File: {file_path}")
    print(f"   ├─ Target Capacity: {capacity or '(Not specified, relying on status marks)'}")
    print(f"   └─ Manual Confirmed Overrides: {manual_confirmed or '(None)'}")

    abs_path = resolve_path(file_path)
    if not os.path.exists(abs_path):
        return (
            f"Error: Registration file not found at '{file_path}'. "
            "Please upload your CSV or Excel file or paste the attendee names directly into the chat."
        )

    ext = os.path.splitext(abs_path.lower())[1]

    raw_rows = []

    # 1. Parse based on file type
    if ext == ".csv":
        print("📊 Reading CSV file...")
        try:
            with open(abs_path, "r", encoding="utf-8-sig") as f:
                reader = csv.reader(f)
                headers = [h.strip() for h in next(reader, [])]
                for r in reader:
                    if r:
                        raw_rows.append(dict(zip(headers, r)))
        except Exception as e:
            # Fallback to alternative encoding if utf-8 fails
            print(f"⚠️ UTF-8 read failed, trying ISO-8859-1 fallback: {e}")
            try:
                with open(abs_path, "r", encoding="ISO-8859-1") as f:
                    reader = csv.reader(f)
                    headers = [h.strip() for h in next(reader, [])]
                    for r in reader:
                        if r:
                            raw_rows.append(dict(zip(headers, r)))
            except Exception as e2:
                return f"Error: Failed to parse CSV file '{file_path}': {e2}"
    elif ext in [".xlsx", ".xls"]:
        print("📊 Reading Excel file via openpyxl...")
        try:
            import openpyxl

            wb = openpyxl.load_workbook(abs_path, data_only=True)
            sheet = wb.active
            rows_iter = sheet.iter_rows(values_only=True)
            headers = [str(h).strip() if h is not None else "" for h in next(rows_iter, [])]
            for r in rows_iter:
                if any(cell is not None for cell in r):
                    row_dict = {}
                    for col_name, val in zip(headers, r):
                        if col_name:
                            # format date values
                            if isinstance(val, (datetime.datetime, datetime.date)):
                                row_dict[col_name] = val.strftime("%Y-%m-%d %H:%M:%S")
                            else:
                                row_dict[col_name] = str(val).strip() if val is not None else ""
                    raw_rows.append(row_dict)
        except Exception as e:
            return f"Error: Failed to parse Excel file '{file_path}': {e}"
    else:
        return f"Error: Unsupported file format '{ext}'! Please upload a .csv or .xlsx file."

    print(f"📊 Loaded {len(raw_rows)} total raw rows from file headers: {headers}")

    # 2. Dynamically identify columns
    first_name_col = ""
    last_name_col = ""
    full_name_col = ""
    timestamp_col = ""
    status_col = ""

    # Normalize headers for matching
    norm_headers = {h.lower().replace("_", "").replace(" ", ""): h for h in headers}

    # Let's map timestamp column with explicit priority
    # Priority 1: Specific registration timestamps
    # Priority 2: Generic creation timestamps
    # Priority 3: Simple dates (excluding check-ins)
    prio1_timestamp = ""
    prio2_timestamp = ""
    prio3_timestamp = ""

    # Column mapping heuristics
    for norm_h, original_h in norm_headers.items():
        # First Name matching
        if norm_h in ["firstname", "name", "givenname", "first"]:
            first_name_col = original_h
        # Last Name matching
        elif norm_h in ["lastname", "surname", "familyname", "last"]:
            last_name_col = original_h
        # Full Name matching
        elif norm_h in ["fullname", "speaker", "fio"]:
            full_name_col = original_h
        # Pre-existing status/waitlist column matching
        elif any(kw in norm_h for kw in ["status", "waitlist", "queue", "state"]):
            status_col = original_h

        # Timestamp/Registration Date matching (prioritized, strictly avoiding check-in columns)
        is_checkin = any(kw in norm_h for kw in ["checkin", "check_in", "check-in", "check"])
        if not is_checkin:
            if any(kw in norm_h for kw in ["registration", "registered", "regdate", "regtime"]):
                prio1_timestamp = original_h
            elif any(kw in norm_h for kw in ["timestamp", "created", "createdat"]):
                prio2_timestamp = original_h
            elif any(kw in norm_h for kw in ["date", "time"]):
                prio3_timestamp = original_h

    timestamp_col = prio1_timestamp or prio2_timestamp or prio3_timestamp

    print("🔍 Dynamic Column Detection Results:")
    print(f"   ├─ First Name Col: {first_name_col or '(not found)'}")
    print(f"   ├─ Last Name Col: {last_name_col or '(not found)'}")
    print(f"   ├─ Full Name Col: {full_name_col or '(not found)'}")
    print(f"   ├─ Status/Waitlist Col: {status_col or '(not found)'}")
    print(f"   └─ Registration Date Col: {timestamp_col or '(not found)'}")

    processed_records = []

    # 3. Clean and map entries
    for row in raw_rows:
        fn = ""
        ln = ""
        reg_date_str = row.get(timestamp_col, "") if timestamp_col else ""

        # Resolve names based on available columns
        if first_name_col and last_name_col:
            fn = row.get(first_name_col, "")
            ln = row.get(last_name_col, "")
        elif full_name_col:
            full = row.get(full_name_col, "").strip()
            parts = full.split(None, 1)
            fn = parts[0] if parts else ""
            ln = parts[1] if len(parts) > 1 else ""
        elif first_name_col:
            # Fallback if only single name column matches First Name criteria
            full = row.get(first_name_col, "").strip()
            parts = full.split(None, 1)
            fn = parts[0] if parts else ""
            ln = parts[1] if len(parts) > 1 else ""

        # Check pre-existing waitlist indicator
        is_waitlisted_pre = False
        if status_col:
            status_val = str(row.get(status_col, "")).strip().lower()
            if any(kw in status_val for kw in ["wait", "pend", "queu"]):
                is_waitlisted_pre = True

        processed_records.append(
            {
                "first_name_raw": fn,
                "last_name_raw": ln,
                "reg_date_str": reg_date_str,
                "parsed_date": parse_date(reg_date_str),
                "is_waitlisted_pre": is_waitlisted_pre,
            }
        )

    # Load GDG Organisers configuration
    organisers_list = load_organisers()
    organisers_lower = {name.lower().strip() for name in organisers_list}

    # 4. Remove duplicate entries based on clean names
    # Sort standard entries by date first to keep the earliest registration
    processed_records.sort(key=lambda x: x["parsed_date"])

    standard_entries_cleaned = []
    seen_names = set()
    duplicate_count = 0

    for rec in processed_records:
        fn_raw = rec["first_name_raw"]
        ln_raw = rec["last_name_raw"]
        fn_clean, ln_clean, invalid_reason = clean_and_validate_name(fn_raw, ln_raw)

        if not fn_clean and not ln_clean:
            continue

        name_key = f"{fn_clean.lower()}|{ln_clean.lower()}"
        if name_key in seen_names:
            duplicate_count += 1
            continue

        seen_names.add(name_key)

        # Check if they are in organisers list
        full_name_clean = f"{fn_clean} {ln_clean}".strip().lower()
        is_org = full_name_clean in organisers_lower

        standard_entries_cleaned.append(
            {
                "first_name": fn_clean,
                "last_name": ln_clean,
                "reg_date": rec["reg_date_str"],
                "parsed_date": rec["parsed_date"],
                "is_waitlisted_pre": rec["is_waitlisted_pre"],
                "invalid_reason": invalid_reason,
                "is_organiser": is_org,
                "is_manual": False,
                "status_override": "organiser" if is_org else "",
            }
        )

    print(f"🧹 Removed {duplicate_count} duplicate registration entries.")

    # 5. Parse manual confirmed override additions
    manual_entries_cleaned = []
    if manual_confirmed:
        # Split by comma or newline
        for part in re.split(r"[,\n]", manual_confirmed):
            part = part.strip()
            if part:
                m_fn, m_ln, status_override = parse_manual_name_line(part)
                if m_fn or m_ln:
                    name_key = f"{m_fn.lower()}|{m_ln.lower()}"
                    full_name_clean = f"{m_fn} {m_ln}".strip().lower()

                    # Forceorganiser classification if matches organisers config list
                    if full_name_clean in organisers_lower or status_override == "organiser":
                        status_override = "organiser"

                    # Check if this manual addition already exists in standard entries
                    found_in_standard = False
                    for s in standard_entries_cleaned:
                        if f"{s['first_name'].lower()}|{s['last_name'].lower()}" == name_key:
                            s["is_manual"] = True
                            if status_override:
                                s["status_override"] = status_override
                            elif s["is_organiser"]:
                                s["status_override"] = "organiser"
                            found_in_standard = True
                            break

                    if not found_in_standard:
                        manual_entries_cleaned.append(
                            {
                                "first_name": m_fn,
                                "last_name": m_ln,
                                "reg_date": "Manual Addition",
                                "parsed_date": datetime.datetime.min,  # Pushes them to the top chronologically
                                "is_waitlisted_pre": False,
                                "invalid_reason": "",
                                "is_organiser": status_override == "organiser",
                                "is_manual": True,
                                "status_override": status_override,
                            }
                        )

    # 6. Separate entries into priority lists
    organisers_tier = []
    speakers_tier = []
    manual_confirmed_tier = []
    standard_valid_entries = []
    waitlist_invalid_entries = []

    # Process all standard and manual unique entries together
    all_unique_entries = standard_entries_cleaned + manual_entries_cleaned

    for entry in all_unique_entries:
        full_name_clean = f"{entry['first_name']} {entry['last_name']}".strip().lower()

        is_org = (
            entry["is_organiser"] or (full_name_clean in organisers_lower) or (entry["status_override"] == "organiser")
        )
        is_spk = entry["status_override"] == "speaker"

        if is_org:
            entry["status_override"] = "organiser"
            organisers_tier.append(entry)
        elif is_spk:
            entry["status_override"] = "speaker"
            speakers_tier.append(entry)
        elif entry["is_manual"]:
            entry["status_override"] = "registered"
            manual_confirmed_tier.append(entry)
        else:
            # Standard valid or invalid
            if entry["invalid_reason"]:
                entry["status_override"] = "waitlist"
                waitlist_invalid_entries.append(entry)
            else:
                standard_valid_entries.append(entry)

    # Sort standard entries chronologically first for capacity limit partitioning
    standard_valid_entries.sort(key=lambda x: x["parsed_date"])

    confirmed_standard = []
    waitlist_standard = []

    if capacity and capacity > 0:
        # Seating capacity IS specified!
        # Max standard slots = capacity - (organisers + speakers + manual additions)
        num_prioritized = len(organisers_tier) + len(speakers_tier) + len(manual_confirmed_tier)
        available_slots = max(0, capacity - num_prioritized)

        for idx, rec in enumerate(standard_valid_entries):
            if idx < available_slots:
                rec["status_override"] = "registered"
                confirmed_standard.append(rec)
            else:
                rec["status_override"] = "waitlist"
                waitlist_standard.append(rec)
    else:
        # Seating capacity is NOT specified! Partition strictly based on status columns inside the file!
        for rec in standard_valid_entries:
            if rec.get("is_waitlisted_pre"):
                rec["status_override"] = "waitlist"
                waitlist_standard.append(rec)
            else:
                rec["status_override"] = "registered"
                confirmed_standard.append(rec)

    # 7. Sort Confirmed List and Waitlist groups separately by script priority (Latin -> Cyrillic -> Other) and alphabetically!
    def sort_group(group_list):
        group_list.sort(
            key=lambda x: (
                get_script_priority(x["first_name"], x["last_name"])[0],
                x["first_name"].lower(),
                x["last_name"].lower(),
            )
        )

    sort_group(organisers_tier)
    sort_group(speakers_tier)
    sort_group(manual_confirmed_tier)
    sort_group(confirmed_standard)
    sort_group(waitlist_standard)
    sort_group(waitlist_invalid_entries)

    # Combine into two distinct lists with correct statuses
    registered_list = []
    for att in organisers_tier:
        registered_list.append({"first_name": att["first_name"], "last_name": att["last_name"], "status": "organiser"})
    for att in speakers_tier:
        registered_list.append({"first_name": att["first_name"], "last_name": att["last_name"], "status": "speaker"})
    for att in manual_confirmed_tier:
        registered_list.append({"first_name": att["first_name"], "last_name": att["last_name"], "status": "registered"})
    for att in confirmed_standard:
        registered_list.append({"first_name": att["first_name"], "last_name": att["last_name"], "status": "registered"})

    waitlist_list = []
    for att in waitlist_standard:
        waitlist_list.append({"first_name": att["first_name"], "last_name": att["last_name"], "status": "waitlist"})
    for att in waitlist_invalid_entries:
        waitlist_list.append({"first_name": att["first_name"], "last_name": att["last_name"], "status": "waitlist"})

    # Assign sequential indexing continuously inside each of the two lists
    for idx, att in enumerate(registered_list):
        att["no"] = idx + 1

    for idx, att in enumerate(waitlist_list):
        att["no"] = idx + 1

    print("📊 Partitioning Complete:")
    print(f"   ├─ Organisers Confirmed: {len(organisers_tier)}")
    print(f"   ├─ Speakers Confirmed: {len(speakers_tier)}")
    # 8. Save results in a beautiful Microsoft Word (.docx) document using 3 columns section layout
    results_dir = os.path.join(BASE_DIR, "results")
    os.makedirs(results_dir, exist_ok=True)

    doc_path = os.path.join(results_dir, "registrations_processed.docx")

    doc_ok = False
    try:
        doc = Document()

        # Helper to set up columns in an existing section
        def set_section_columns(section, num_cols: int, space_twips: int = 288):
            sectPr = section._sectPr
            cols = sectPr.find(qn("w:cols"))
            if cols is not None:
                cols.set(qn("w:num"), str(num_cols))
                cols.set(qn("w:space"), str(space_twips))
            else:
                cols = OxmlElement("w:cols")
                cols.set(qn("w:num"), str(num_cols))
                cols.set(qn("w:space"), str(space_twips))
                sectPr.append(cols)

        # Configure layout margins and 3 columns on the main/only section
        sec_tables = doc.sections[0]
        sec_tables.top_margin = Inches(0.4)
        sec_tables.bottom_margin = Inches(0.4)
        sec_tables.left_margin = Inches(0.5)
        sec_tables.right_margin = Inches(0.5)
        set_section_columns(sec_tables, 3, space_twips=288)  # 3 columns with 0.2 inch spacing

        # Helper to set background shading in table cells
        def set_cell_background(cell, fill_hex: str):
            shading_xml = f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'
            cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

        # Helper function to render a participant table inside the current section column
        def build_docx_table(attendees, title_text, header_color_hex):
            # Table Title
            h = doc.add_paragraph()
            h_run = h.add_run(title_text)
            h_run.font.name = "Arial"
            h_run.font.size = Pt(10.5)
            h_run.font.bold = True
            h_run.font.color.rgb = RGBColor(
                int(header_color_hex[0:2], 16), int(header_color_hex[2:4], 16), int(header_color_hex[4:6], 16)
            )
            h.paragraph_format.space_before = Pt(6)
            h.paragraph_format.space_after = Pt(3)

            if not attendees:
                p = doc.add_paragraph()
                p_run = p.add_run("*List is empty*")
                p_run.font.name = "Arial"
                p_run.font.size = Pt(8.5)
                p_run.font.italic = True
                return

            table = doc.add_table(rows=0, cols=3)
            table.allow_autofit = False

            # Setup columns widths (Sum equals exactly 2.3 inches, fitting 3-column layout width perfectly!)
            # No. (0.2 in), Full Name (1.55 in), Status (0.55 in)
            col_widths = [0.2, 1.55, 0.55]

            # Formulate Data Rows (Alternate shading white / light gray)
            for r_idx, att in enumerate(attendees):
                row_cells = table.add_row().cells
                fill_color = "F8F9FA" if r_idx % 2 == 1 else "FFFFFF"

                # 1. Write and style No. cell (centered)
                cell0 = row_cells[0]
                set_cell_background(cell0, fill_color)
                p0 = cell0.paragraphs[0]
                p0.text = ""
                r0 = p0.add_run(str(att["no"]))
                r0.font.name = "Arial"
                r0.font.size = Pt(8.0)
                p0.paragraph_format.alignment = 1  # Center
                p0.paragraph_format.space_before = Pt(1.5)
                p0.paragraph_format.space_after = Pt(1.5)

                # 2. Write and style Full Name cell (left aligned)
                cell1 = row_cells[1]
                set_cell_background(cell1, fill_color)
                p1 = cell1.paragraphs[0]
                p1.text = ""

                raw_name = f"{att['first_name']} {att['last_name']}".strip()
                clean_name = re.sub(r"\d+", "", raw_name)
                clean_name = re.sub(r"\s+", " ", clean_name).strip()

                r1 = p1.add_run(clean_name)
                r1.font.name = "Arial"
                r1.font.size = Pt(8.0)
                p1.paragraph_format.alignment = 0  # Left
                p1.paragraph_format.space_before = Pt(1.5)
                p1.paragraph_format.space_after = Pt(1.5)

                # 3. Write and style Status cell (centered, custom colors)
                cell2 = row_cells[2]
                set_cell_background(cell2, fill_color)
                p2 = cell2.paragraphs[0]
                p2.text = ""

                status_text = att["status"]
                if status_text in ["registered", "waitlist"]:
                    status_text = ""

                r2 = p2.add_run(status_text)
                r2.font.name = "Arial"
                r2.font.size = Pt(7.5)
                r2.font.bold = False
                if att["status"] == "organiser":
                    r2.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)  # Google Blue
                elif att["status"] == "speaker":
                    r2.font.color.rgb = RGBColor(0x0F, 0x9D, 0x58)  # Google Green

                p2.paragraph_format.alignment = 1  # Center
                p2.paragraph_format.space_before = Pt(1.5)
                p2.paragraph_format.space_after = Pt(1.5)

            # Apply widths to all columns explicitly
            for idx, width in enumerate(col_widths):
                table.columns[idx].width = Inches(width)

            # Apply widths to all rows cells
            for row in table.rows:
                for idx, width in enumerate(col_widths):
                    row.cells[idx].width = Inches(width)

            # Spacer below table
            doc.add_paragraph().paragraph_format.space_before = Pt(3)

        # Draw Registered list table
        build_docx_table(registered_list, "🟢 Confirmed Registrations", "1A73E8")

        # Draw Waitlist table
        build_docx_table(waitlist_list, "🔴 Waitlist", "EA4335")

        doc.save(doc_path)
        doc_ok = True
        print(f"✅ Successfully wrote beautiful DOCX document to: {doc_path}")
    except Exception as docx_err:
        print(f"⚠️ Failed to generate DOCX document: {docx_err}")

    # 9. Simple, professional chat summary containing only the statistics and file links (no large inline tables)
    summary_str = f"""### 📊 Event Registration Processing Summary
Successfully processed the registrations list:
- **Total Raw Entries**: {len(raw_rows)}
- **Duplicates Removed**: {duplicate_count}
- **Confirmed Registrants (registered)**: **{len(registered_list)}** (including {len(manual_entries_cleaned)} manual additions)
  * *Organisers*: {len(organisers_tier)}
  * *Speakers*: {len(speakers_tier)}
  * *Registered Attendees*: {len(manual_confirmed_tier) + len(confirmed_standard)}
- **Waitlisted Registrants (waitlist)**: **{len(waitlist_list)}**

---

### 📂 Generated Deliverables:
"""
    if doc_ok:
        summary_str += f"- **Premium Event Registration Document (DOCX, 3 Columns)**: [{os.path.basename(doc_path)}](file://{doc_path})\n"
    else:
        summary_str += "⚠️ Failed to generate Word document due to an error.\n"
    return summary_str


# ---------------------------------------------------------------------------
# Tools for Managing Organisers List Dynamically
# ---------------------------------------------------------------------------


def _get_organisers_file_path() -> str:
    """Helper to locate or initialize configs/organisers.txt."""
    parent_dir = os.path.dirname(BASE_DIR)
    project_root = os.path.dirname(parent_dir)

    for p in [
        os.path.join(project_root, "configs", "organisers.txt"),
        os.path.join(parent_dir, "configs", "organisers.txt"),
    ]:
        if os.path.exists(p):
            return p

    local_path = os.path.join(BASE_DIR, "organisers.txt")
    if os.path.exists(local_path):
        return local_path

    # If none exists, create in project root configs/
    configs_dir = os.path.join(project_root, "configs")
    os.makedirs(configs_dir, exist_ok=True)
    return os.path.join(configs_dir, "organisers.txt")


def get_organisers_list() -> dict:
    """
    Reads and returns the list of official organizers from the organisers.txt configuration file.

    Returns:
        dict: Success status, the list of organizer names, and file path.
    """
    path = _get_organisers_file_path()
    try:
        if not os.path.exists(path):
            default_orgs = [
                "Aliaksandr Aplachykau",
                "Yuliya Aplachykava",
                "Aryna Stsiapanava",
                "Roman Hordiichuk",
                "Paul Kastel",
                "Rafał Płonka",
                "Yusuf Gültaç",
            ]
            os.makedirs(os.path.dirname(path), exist_ok=True)
            with open(path, "w", encoding="utf-8") as f:
                f.write("\n".join(default_orgs) + "\n")
            return {"success": True, "organisers": default_orgs, "file_path": path}

        with open(path, "r", encoding="utf-8") as f:
            orgs = [line.strip() for line in f if line.strip()]
        return {"success": True, "organisers": orgs, "file_path": path}
    except Exception as e:
        return {"success": False, "error": f"Failed to read organisers: {str(e)}"}


def add_organiser(name: str) -> dict:
    """
    Adds a new organizer to the official organisers.txt configuration file.
    Automatically normalizes and formats the name (capitalizing each word).

    Args:
        name: Full name of the organizer to add (e.g. "John Doe").

    Returns:
        dict: Success status, name added, and custom response message.
    """
    path = _get_organisers_file_path()
    parts = name.strip().split()
    if not parts:
        return {"success": False, "error": "Invalid name provided."}
    cleaned_name = " ".join(part.capitalize() for part in parts)

    try:
        res = get_organisers_list()
        if not res.get("success"):
            return res
        orgs = res["organisers"]

        orgs_lower = [o.lower() for o in orgs]
        if cleaned_name.lower() in orgs_lower:
            return {
                "success": True,
                "name_added": cleaned_name,
                "message": f"Organiser '{cleaned_name}' is already in the list.",
            }

        orgs.append(cleaned_name)
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(orgs) + "\n")

        return {
            "success": True,
            "name_added": cleaned_name,
            "message": f"Successfully added '{cleaned_name}' to organisers.txt.",
        }
    except Exception as e:
        return {"success": False, "error": f"Failed to add organiser: {str(e)}"}


def remove_organiser(name: str) -> dict:
    """
    Removes an organizer from the official organisers.txt configuration file.
    Intelligently performs approximate case-insensitive matching to find the organizer.

    Args:
        name: Full name or partial name of the organizer to remove.

    Returns:
        dict: Success status, name removed, and custom response message.
    """
    path = _get_organisers_file_path()
    target = name.strip().lower()
    if not target:
        return {"success": False, "error": "Invalid name provided."}

    try:
        res = get_organisers_list()
        if not res.get("success"):
            return res
        orgs = res["organisers"]

        matched_name = None
        for o in orgs:
            if o.lower() == target:
                matched_name = o
                break

        if not matched_name:
            matches = []
            for o in orgs:
                o_lower = o.lower()
                if target in o_lower or any(part in o_lower for part in target.split()):
                    matches.append(o)
            if len(matches) == 1:
                matched_name = matches[0]
            elif len(matches) > 1:
                return {
                    "success": False,
                    "error": f"Multiple potential organisers found: {', '.join(matches)}. Please be more specific.",
                }

        if not matched_name:
            return {"success": False, "error": f"No organiser matching '{name}' was found in organisers.txt."}

        orgs.remove(matched_name)
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(orgs) + "\n")

        return {
            "success": True,
            "name_removed": matched_name,
            "message": f"Successfully removed '{matched_name}' from organisers.txt.",
        }
    except Exception as e:
        return {"success": False, "error": f"Failed to remove organiser: {str(e)}"}
